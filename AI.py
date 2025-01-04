import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import customtkinter as ctk
import time
import random
import json
from docx import Document
import os
import threading
import requests
from textblob import TextBlob
import subprocess
import re
from utils import load_api_settings
import nltk
import os

import nltk
nltk.download('punkt_tab')


API_URL = "https://api-inference.huggingface.co/models/facebook/bart-large-cnn"
API_TOKEN = "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"

headers = {"Authorization": f"Bearer {API_TOKEN}"}

def segment_text(text, max_tokens=500):
    words = text.split()
    segments = []
    for i in range(0, len(words), max_tokens):
        segment = ' '.join(words[i:i+max_tokens])
        segments.append(segment)
    return segments

def query(payload, max_retries=5):
    for attempt in range(max_retries):
        try:
            response = requests.post(API_URL, headers=headers, json=payload, timeout=30)
            response.raise_for_status()
            return response.json()
        except requests.exceptions.RequestException as e:
            if attempt == max_retries - 1:
                raise Exception(f"Failed after {max_retries} attempts: {str(e)}")
            time.sleep(2 ** attempt + random.uniform(0, 1))
    raise Exception(f"Failed after {max_retries} attempts")

class AISummaryWidget(ctk.CTkToplevel):
    def __init__(self, parent, news_data):
        super().__init__(parent)
        self.title("AI Financial News Analysis")
        self.geometry("600x700")
        self.minsize(600, 700)
        self.news_data = news_data
        self.analysis_result = ""
        self.parent = parent
        self.api_settings = load_api_settings()
        self.selected_platform = tk.StringVar(value="No API Selected")
        self.selected_api_name = tk.StringVar(value="")
        self.selected_api_key = tk.StringVar(value="")
        
        self.load_selected_api()
        self.create_widgets()
        self.update_generate_button_state()  # Move this here
        
        self.after(10, self.center_window)
        
        self.transient(parent)
        self.grab_set()
        self.protocol("WM_DELETE_WINDOW", self.on_closing)


    def load_selected_api(self):
        try:
            with open('selected_api.json', 'r') as f:
                selected = json.load(f)
                name = selected['name']
                if name in self.api_settings:
                    details = self.api_settings[name]
                    self.selected_platform.set(details['platform'])
                    self.selected_api_name.set(name)
                    masked_token = details['token'][:5] + "*" * (len(details['token']) - 5)
                    self.selected_api_key.set(masked_token)
        except FileNotFoundError:
            pass  # No previously selected API
            
    def save_selected_api(self, name, platform, token):
        selected = {'name': name}
        with open('selected_api.json', 'w') as f:
            json.dump(selected, f)

    def center_window(self):
        self.update_idletasks()
        width = self.winfo_width()
        height = self.winfo_height()
        screen_width = self.winfo_screenwidth()
        screen_height = self.winfo_screenheight()
        
        x = (screen_width - width) // 2
        y = (screen_height - height) // 2
        
        self.geometry(f'{width}x{height}+{x}+{y}')
        
        # Ensure the window is on top and focused
        self.lift()
        self.focus_force()
        
    def create_widgets(self):
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(1, weight=1)

        # API Info Frame
        api_frame = ctk.CTkFrame(self)
        api_frame.grid(row=0, column=0, sticky="ew", padx=10, pady=10)
        api_frame.grid_columnconfigure(1, weight=1)

        ctk.CTkLabel(api_frame, text="Selected API:").grid(row=0, column=0, padx=5, pady=5, sticky="w")
        ctk.CTkLabel(api_frame, textvariable=self.selected_platform).grid(row=0, column=1, padx=5, pady=5, sticky="w")
        
        ctk.CTkLabel(api_frame, text="API Name:").grid(row=1, column=0, padx=5, pady=5, sticky="w")
        ctk.CTkLabel(api_frame, textvariable=self.selected_api_name).grid(row=1, column=1, padx=5, pady=5, sticky="w")
        
        ctk.CTkLabel(api_frame, text="API Key:").grid(row=2, column=0, padx=5, pady=5, sticky="w")
        ctk.CTkLabel(api_frame, textvariable=self.selected_api_key).grid(row=2, column=1, padx=5, pady=5, sticky="w")

        ctk.CTkButton(api_frame, text="Select API", command=self.show_select_dialog).grid(row=3, column=0, columnspan=2, pady=5)


        # Main Content Frame
        main_frame = ctk.CTkFrame(self)
        main_frame.grid(row=1, column=0, sticky="nsew", padx=10, pady=10)
        main_frame.grid_columnconfigure(0, weight=1)
        main_frame.grid_rowconfigure(0, weight=1)

        self.summary_text = ctk.CTkTextbox(main_frame, wrap=tk.WORD)
        self.summary_text.grid(row=0, column=0, sticky="nsew", padx=10, pady=10)
        
        button_frame = ctk.CTkFrame(main_frame)
        button_frame.grid(row=1, column=0, sticky="ew", padx=10, pady=10)
        button_frame.grid_columnconfigure((0, 1), weight=1)
        
        self.generate_button = ctk.CTkButton(button_frame, text="Generate Analysis", command=self.start_analysis)
        self.generate_button.grid(row=0, column=0, padx=5, pady=5, sticky="ew")
                
        # Disable the generate button if no API is selected
        self.update_generate_button_state()
        
        self.export_button = ctk.CTkButton(button_frame, text="Export to Word", command=self.export_to_word)
        self.export_button.grid(row=0, column=1, padx=5, pady=5, sticky="ew")
        self.export_button.configure(state=tk.DISABLED)
        
        self.progress_frame = ctk.CTkFrame(main_frame)
        self.progress_frame.grid(row=2, column=0, sticky="ew", padx=10, pady=10)
        self.progress_frame.grid_columnconfigure(0, weight=1)
        
        self.progress_label = ctk.CTkLabel(self.progress_frame, text="")
        self.progress_label.grid(row=0, column=0, sticky="w", padx=5, pady=2)
        
        self.progress_bar = ctk.CTkProgressBar(self.progress_frame)
        self.progress_bar.grid(row=1, column=0, sticky="ew", padx=5, pady=2)
        self.progress_bar.set(0)

    def load_api_settings(self):
        try:
            with open('api_settings.json', 'r') as f:
                return json.load(f)
        except FileNotFoundError:
            return {"AI Community Platform": {"platform": "AI Community Platform", "token": ""}, 
                    "OpenAI ChatGPT": {"platform": "OpenAI ChatGPT", "token": ""}}

    def save_api_settings(self):
        with open('api_settings.json', 'w') as f:
            json.dump(self.api_settings, f)

    def show_select_dialog(self):
        self.grab_release()  # Release the grab before opening the dialog
        dialog = APISelectDialog(self, self.api_settings)
        self.wait_window(dialog)
        if dialog.selected_api:
            name, platform, token = dialog.selected_api
            self.selected_platform.set(platform)
            self.selected_api_name.set(name)  # Set the API name
            masked_token = token[:5] + "*" * (len(token) - 5)
            self.selected_api_key.set(masked_token)
            self.api_settings = dialog.api_settings 
            self.save_selected_api(name, platform, token)
            messagebox.showinfo("API Selected", f"Selected API: {name} ({platform})")
        self.grab_set()  # Re-establish the grab after the dialog is closed
        self.focus_force()  # Force focus back to this window
        self.lift()  # Lift the window to the top
        self.update_generate_button_state()

    def update_generate_button_state(self):
        if self.selected_platform.get() != "No API Selected" and self.selected_api_key.get():
            self.generate_button.configure(state=tk.NORMAL)
        else:
            self.generate_button.configure(state=tk.DISABLED)

    def on_closing(self):
        self.grab_release()
        self.destroy()

    def on_platform_change(self, choice):
        self.api_token_entry.delete(0, tk.END)
        self.api_token_entry.insert(0, self.api_settings.get(choice, ""))

    def add_api_token(self):
        platform = self.platform_var.get()
        token = self.api_token_entry.get()
        if token:
            self.api_settings[platform] = token
            self.save_api_settings()
            messagebox.showinfo("Success", f"API token for {platform} has been added.")
        else:
            messagebox.showwarning("Invalid Input", "Please enter an API token.")

    def edit_api_token(self):
        platform = self.platform_var.get()
        token = self.api_token_entry.get()
        if token:
            self.api_settings[platform] = token
            self.save_api_settings()
            messagebox.showinfo("Success", f"API token for {platform} has been updated.")
        else:
            messagebox.showwarning("Invalid Input", "Please enter an API token.")

    def select_api_token(self):
        platform = self.platform_var.get()
        token = self.api_settings.get(platform, "")
        self.api_token_entry.delete(0, tk.END)
        self.api_token_entry.insert(0, token)

    def test_api_connection(self):
        platform = self.platform_var.get()
        token = self.api_token_entry.get()
        
        if not token:
            messagebox.showwarning("Invalid Input", "Please enter an API token.")
            return

        try:
            if platform == "AI Community Platform":
                API_URL = "https://api-inference.huggingface.co/models/facebook/bart-large-cnn"
                headers = {"Authorization": f"Bearer {token}"}
                response = requests.post(API_URL, headers=headers, json={"inputs": "Hello, world!"})
                response.raise_for_status()
            elif platform == "OpenAI ChatGPT":
                API_URL = "https://api.openai.com/v1/chat/completions"
                headers = {"Authorization": f"Bearer {token}"}
                data = {
                    "model": "gpt-3.5-turbo",
                    "messages": [{"role": "user", "content": "Hello, world!"}],
                    "max_tokens": 5
                }
                response = requests.post(API_URL, headers=headers, json=data)
                response.raise_for_status()
            
            messagebox.showinfo("Success", f"Connection to {platform} API successful!")
        except requests.exceptions.RequestException as e:
            messagebox.showerror("Connection Error", f"Failed to connect to {platform} API: {str(e)}")
            
            
    
    def start_analysis(self):
        if self.selected_platform.get() == "No API Selected" or not self.selected_api_key.get():
            messagebox.showwarning("No API Selected", "Please select a valid API before generating analysis.")
            return

        self.generate_button.configure(state=tk.DISABLED)
        self.export_button.configure(state=tk.DISABLED)
        self.summary_text.delete('1.0', tk.END)
        self.progress_bar.set(0)
        self.progress_label.configure(text="Initializing analysis...")
        
        threading.Thread(target=self.generate_analysis, daemon=True).start()
    

    def generate_analysis(self):
        try:
            platform = self.selected_platform.get()
            name = self.selected_api_name.get()
            token = self.api_settings[name]['token']  # Get the actual token
            self.analysis_result = self.perform_analysis()  # Remove arguments here
            self.after(0, self.update_ui_with_result)
        except Exception as e:
            error_message = f"Failed to generate analysis: {str(e)}"
            self.after(0, lambda: messagebox.showerror("Error", error_message))
        finally:
            self.after(0, self.reset_ui)
    
    def perform_analysis(self):
        platform = self.selected_platform.get()
        analysis = f"AI Financial News Analysis using {platform}\n\n"

        # Implement your API call here using the token if needed
        # For now, we'll just use the existing methods
        analysis += self.analyze_market_trends()
        analysis += self.identify_stocks()
        analysis += self.analyze_malaysia_market()
        analysis += self.analyze_us_market()
        analysis += self.analyze_overall_sentiment()
        analysis += self.list_market_factors()

        return analysis


    def analyze_market_trends(self):
        trends = "1. Key Market Trends\n"
        for news in self.news_data:
            if 'market' in news['title'].lower() or 'trend' in news['title'].lower():
                trends += f"- {news['title']}\n"
        return trends + "\n" if trends != "1. Key Market Trends\n" else "1. Key Market Trends\nNo specific market trends identified in the news data.\n\n"

    def identify_stocks(self):
        stocks = "2. Stocks Mentioned\n"
        stock_pattern = r'\b[A-Z]{1,5}\b'  # Pattern to match stock symbols
        mentioned_stocks = set()
        for news in self.news_data:
            matches = re.findall(stock_pattern, news['title'] + " " + news['description'])
            mentioned_stocks.update(matches)
        if mentioned_stocks:
            stocks += ", ".join(mentioned_stocks) + "\n"
        return stocks + "\n" if stocks != "2. Stocks Mentioned\n" else "2. Stocks Mentioned\nNo specific stocks identified in the news data.\n\n"

    def analyze_malaysia_market(self):
        analysis = "3. Malaysia Market Analysis\n"
        malaysia_news = [news for news in self.news_data if 'malaysia' in news['title'].lower() or 'malaysia' in news['description'].lower()]
        if malaysia_news:
            sentiments = [float(news['sentiment']) for news in malaysia_news]
            avg_sentiment = sum(sentiments) / len(sentiments)
            analysis += f"- Short-term outlook (1-3 months): {'Positive' if avg_sentiment > 0 else 'Negative'}\n"
            analysis += f"- Long-term outlook (6-12 months): {'Stable to Positive' if avg_sentiment > -0.2 else 'Uncertain'}\n"
            analysis += f"- Market sentiment score: {5 + int(avg_sentiment * 5)}/10\n"
            analysis += "- Notable news:\n"
            for news in malaysia_news[:3]:  # List top 3 news items
                analysis += f"  * {news['title']}\n"
        return analysis + "\n" if analysis != "3. Malaysia Market Analysis\n" else "3. Malaysia Market Analysis\nNo specific information available for the Malaysia market.\n\n"

    def analyze_us_market(self):
        analysis = "4. US Market Analysis\n"
        us_news = [news for news in self.news_data if 'us' in news['title'].lower() or 'us' in news['description'].lower()]
        if us_news:
            sentiments = [float(news['sentiment']) for news in us_news]
            avg_sentiment = sum(sentiments) / len(sentiments)
            analysis += f"- Short-term outlook (1-3 months): {'Positive' if avg_sentiment > 0 else 'Negative'}\n"
            analysis += f"- Long-term outlook (6-12 months): {'Stable to Positive' if avg_sentiment > -0.2 else 'Uncertain'}\n"
            analysis += f"- Market sentiment score: {5 + int(avg_sentiment * 5)}/10\n"
            analysis += "- Notable news:\n"
            for news in us_news[:3]:  # List top 3 news items
                analysis += f"  * {news['title']}\n"
        return analysis + "\n" if analysis != "4. US Market Analysis\n" else "4. US Market Analysis\nNo specific information available for the US market.\n\n"

    def analyze_overall_sentiment(self):
        analysis = "5. Overall Market Sentiment\n"
        sentiments = [float(news['sentiment']) for news in self.news_data]
        avg_sentiment = sum(sentiments) / len(sentiments)
        analysis += f"The overall market sentiment is {'positive' if avg_sentiment > 0 else 'negative'} "
        analysis += f"with an average sentiment score of {avg_sentiment:.2f}.\n"
        analysis += "Key factors influencing markets:\n"
        # Extract key factors (this is a simplified approach)
        factors = set()
        for news in self.news_data:
            blob = TextBlob(news['description'])
            factors.update(blob.noun_phrases)
        analysis += "- " + "\n- ".join(list(factors)[:5])  # List top 5 factors
        return analysis + "\n\n"

    def list_market_factors(self):
        analysis = "6. Positive Factors for Malaysia\n"
        analysis += self.extract_factors('malaysia', positive=True)
        analysis += "7. Negative Factors for Malaysia\n"
        analysis += self.extract_factors('malaysia', positive=False)
        analysis += "8. Positive Factors for US\n"
        analysis += self.extract_factors('us', positive=True)
        analysis += "9. Negative Factors for US\n"
        analysis += self.extract_factors('us', positive=False)
        return analysis

    def extract_factors(self, country, positive=True):
        relevant_news = [news for news in self.news_data if country in news['title'].lower() or country in news['description'].lower()]
        factors = []
        for news in relevant_news:
            sentiment = float(news['sentiment'])
            if (positive and sentiment > 0.1) or (not positive and sentiment < -0.1):
                blob = TextBlob(news['description'])
                factors.extend(blob.noun_phrases)
        
        # Remove duplicates and limit to top 5
        factors = list(set(factors))[:5]
        return "- " + "\n- ".join(factors) + "\n\n" if factors else "No specific factors identified.\n\n"

    def update_progress(self, progress, text):
        self.after(0, lambda: self.progress_bar.set(progress))
        self.after(0, lambda: self.progress_label.configure(text=text))
    
    def update_ui_with_result(self):
        self.summary_text.delete('1.0', tk.END)
        self.summary_text.insert(tk.END, self.analysis_result)
        self.export_button.configure(state=tk.NORMAL)
    
    def reset_ui(self):
        self.generate_button.configure(state=tk.NORMAL)
        self.progress_label.configure(text="Analysis complete")
        self.progress_bar.set(1)
    
    def export_to_word(self):
        file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
        if file_path:
            doc = Document()
            doc.add_heading('AI Financial News Analysis', 0)
            for line in self.analysis_result.split('\n'):
                if line.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
                    doc.add_heading(line, level=1)
                elif line.startswith('-'):
                    doc.add_paragraph(line, style='List Bullet')
                else:
                    doc.add_paragraph(line)
            doc.save(file_path)
            
            if messagebox.askyesno("Export Successful", f"Analysis exported to {file_path}\n\nDo you want to open the file?"):
                os.startfile(file_path)
                
class APISelectDialog(ctk.CTkToplevel):
    def __init__(self, parent, api_settings):
        super().__init__(parent)
        self.title("Select API")
        self.geometry("800x400")
        self.api_settings = api_settings
        self.selected_api = None
        self.platforms = ["AI Community Platform", "OpenAI ChatGPT"]  # Add more platforms as needed
        
        # Make the dialog modal
        self.transient(parent)
        self.grab_set()
        
        # Center the dialog on the parent window
        self.update_idletasks()
        parent_x = parent.winfo_x()
        parent_y = parent.winfo_y()
        parent_width = parent.winfo_width()
        parent_height = parent.winfo_height()
        dialog_width = self.winfo_width()
        dialog_height = self.winfo_height()
        x = parent_x + (parent_width - dialog_width) // 2
        y = parent_y + (parent_height - dialog_height) // 2
        self.geometry(f"+{x}+{y}")
        
        self.create_widgets()


    def create_widgets(self):
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(0, weight=1)

        main_frame = ctk.CTkFrame(self)
        main_frame.grid(row=0, column=0, sticky="nsew", padx=10, pady=10)
        main_frame.grid_columnconfigure(0, weight=1)
        main_frame.grid_rowconfigure(0, weight=1)

        # Create Treeview
        self.tree = ttk.Treeview(main_frame, columns=("Name", "Source", "API Key"), show="headings")
        self.tree.heading("Name", text="Name")
        self.tree.heading("Source", text="Source")
        self.tree.heading("API Key", text="API Key")
        self.tree.column("Name", width=200)
        self.tree.column("Source", width=200)
        self.tree.column("API Key", width=300)
        self.tree.grid(row=0, column=0, sticky="nsew", padx=5, pady=5)


        # Add scrollbar
        scrollbar = ttk.Scrollbar(main_frame, orient="vertical", command=self.tree.yview)
        scrollbar.grid(row=0, column=1, sticky="ns")
        self.tree.configure(yscrollcommand=scrollbar.set)

        self.update_treeview()

        button_frame = ctk.CTkFrame(main_frame)
        button_frame.grid(row=1, column=0, sticky="ew", pady=10)
        button_frame.grid_columnconfigure((0, 1, 2, 3, 4), weight=1)

        ctk.CTkButton(button_frame, text="Add", command=self.add_api).grid(row=0, column=0, padx=5)
        ctk.CTkButton(button_frame, text="Edit", command=self.edit_api).grid(row=0, column=1, padx=5)
        ctk.CTkButton(button_frame, text="Delete", command=self.delete_api).grid(row=0, column=2, padx=5)
        ctk.CTkButton(button_frame, text="Test", command=self.test_api).grid(row=0, column=3, padx=5)
        ctk.CTkButton(button_frame, text="Confirm", command=self.confirm_selection).grid(row=0, column=4, padx=5)

    def update_treeview(self):
        self.tree.delete(*self.tree.get_children())
        for name, details in self.api_settings.items():
            masked_token = details['token'][:5] + "*" * (len(details['token']) - 5) if details['token'] else ""
            self.tree.insert("", "end", values=(name, details['platform'], masked_token))

    def update_listbox(self):
        self.listbox.delete(0, tk.END)
        for name, details in self.api_settings.items():
            if isinstance(details, dict):
                masked_token = details['token'][:5] + "*" * (len(details['token']) - 5)
                platform = details['platform']
            else:
                masked_token = details[:5] + "*" * (len(details) - 5)
                platform = name
            self.listbox.insert(tk.END, f"{name} ({platform}): {masked_token}")
            
    def add_api(self):
        add_window = ctk.CTkToplevel(self)
        add_window.title("Add API")
        add_window.geometry("300x400")
        add_window.transient(self)
        add_window.grab_set()

        ctk.CTkLabel(add_window, text="Name:").pack(pady=5)
        name_entry = ctk.CTkEntry(add_window)
        name_entry.pack(pady=5)

        ctk.CTkLabel(add_window, text="Platform:").pack(pady=5)
        platform_var = tk.StringVar(value=self.platforms[0])
        platform_dropdown = ctk.CTkOptionMenu(add_window, variable=platform_var, values=self.platforms)
        platform_dropdown.pack(pady=5)

        ctk.CTkLabel(add_window, text="API Token:").pack(pady=5)
        token_entry = ctk.CTkEntry(add_window, show="*")
        token_entry.pack(pady=5)

        def save_api():
            name = name_entry.get()
            platform = platform_var.get()
            token = token_entry.get()
            if name and token:
                self.api_settings[name] = {'platform': platform, 'token': token}
                self.update_treeview()
                self.save_api_settings()
                add_window.destroy()
            else:
                messagebox.showwarning("Invalid Input", "Please enter both name and API token.")

        ctk.CTkButton(add_window, text="Save", command=save_api).pack(pady=10)

    def edit_api(self):
        selected = self.tree.selection()
        if not selected:
            messagebox.showwarning("No Selection", "Please select an API to edit.")
            return

        item = self.tree.item(selected[0])
        name = item['values'][0]
        details = self.api_settings[name]

        edit_window = ctk.CTkToplevel(self)
        edit_window.title(f"Edit API: {name}")
        edit_window.geometry("300x400")
        edit_window.transient(self)
        edit_window.grab_set()

        ctk.CTkLabel(edit_window, text="Name:").pack(pady=5)
        name_entry = ctk.CTkEntry(edit_window)
        name_entry.insert(0, name)
        name_entry.pack(pady=5)

        ctk.CTkLabel(edit_window, text="Platform:").pack(pady=5)
        platform_var = tk.StringVar(value=details['platform'])
        platform_dropdown = ctk.CTkOptionMenu(edit_window, variable=platform_var, values=self.platforms)
        platform_dropdown.pack(pady=5)

        ctk.CTkLabel(edit_window, text="API Token:").pack(pady=5)
        token_entry = ctk.CTkEntry(edit_window, show="*")
        token_entry.insert(0, details['token'])
        token_entry.pack(pady=5)

        def save_changes():
            new_name = name_entry.get()
            new_platform = platform_var.get()
            new_token = token_entry.get()
            if new_name and new_token:
                if new_name != name:
                    del self.api_settings[name]
                self.api_settings[new_name] = {'platform': new_platform, 'token': new_token}
                self.update_treeview()
                self.save_api_settings()
                edit_window.destroy()
            else:
                messagebox.showwarning("Invalid Input", "Please enter both name and API token.")

        ctk.CTkButton(edit_window, text="Save Changes", command=save_changes).pack(pady=10)

    def delete_api(self):
        selected = self.tree.selection()
        if not selected:
            messagebox.showwarning("No Selection", "Please select an API to delete.")
            return

        item = self.tree.item(selected[0])
        name = item['values'][0]
        if messagebox.askyesno("Confirm Deletion", f"Are you sure you want to delete the API: {name}?"):
            del self.api_settings[name]
            self.update_treeview()
            self.save_api_settings()  # Make sure to save the changes
            messagebox.showinfo("Deletion Successful", f"API '{name}' has been deleted.")

    def test_api(self):
        selected = self.tree.selection()
        if not selected:
            messagebox.showwarning("No Selection", "Please select an API to test.")
            return

        item = self.tree.item(selected[0])
        name = item['values'][0]
        details = self.api_settings[name]
        platform = details['platform']
        token = details['token']
        
        try:
            if platform == "AI Community Platform":
                API_URL = "https://api-inference.huggingface.co/models/gpt2"
                headers = {"Authorization": f"Bearer {token}"}
                response = requests.post(API_URL, headers=headers, json={"inputs": "Hello, world!"})
                response.raise_for_status()
                messagebox.showinfo("API Test", f"{name} ({platform}) API test successful!")
            elif platform == "OpenAI ChatGPT":
                API_URL = "https://api.openai.com/v1/chat/completions"
                headers = {"Authorization": f"Bearer {token}"}
                data = {
                    "model": "gpt-3.5-turbo",
                    "messages": [{"role": "user", "content": "Hello, world!"}],
                    "max_tokens": 5
                }
                response = requests.post(API_URL, headers=headers, json=data)
                response.raise_for_status()
                messagebox.showinfo("API Test", f"{name} ({platform}) API test successful!")
            else:
                messagebox.showwarning("API Test", f"Test not implemented for {platform}")
        except requests.exceptions.RequestException as e:
            messagebox.showerror("API Test Failed", f"Error testing {name} ({platform}) API: {str(e)}")


    def confirm_selection(self):
        selected = self.tree.selection()
        if not selected:
            messagebox.showwarning("No Selection", "Please select an API from the list.")
            return

        item = self.tree.item(selected[0])
        name = item['values'][0]
        details = self.api_settings[name]
        self.selected_api = (name, details['platform'], details['token'])
        self.destroy()
              
    def save_api_settings(self):
        with open('api_settings.json', 'w') as f:
            json.dump(self.api_settings, f, indent=2)             
              

class ProgressWindow(ctk.CTkToplevel):
    def __init__(self, parent):
        super().__init__(parent)
        self.title("Generating Analysis")
        self.geometry("300x100")
        self.transient(parent)
        self.grab_set()

        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(1, weight=1)

        self.label = ctk.CTkLabel(self, text="Initializing...")
        self.label.grid(row=0, column=0, padx=10, pady=5, sticky="ew")

        self.progress = ctk.CTkProgressBar(self)
        self.progress.grid(row=1, column=0, padx=10, pady=10, sticky="ew")
        self.progress.set(0)

    def update_progress(self, value, text):
        self.progress.set(value)
        self.label.configure(text=text)
        self.update()


def create_ai_summary_widget(parent, news_data):
    ai_widget = AISummaryWidget(parent, news_data)
    parent.wait_window(ai_widget)
    return ai_widget.analysis_result